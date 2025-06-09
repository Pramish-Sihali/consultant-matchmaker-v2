# app/services/cv_processor.py - Enhanced CV Processor with Phased Processing

import asyncio
import io
import logging
import re
import time
from typing import Dict, Any, Optional, Callable, Tuple
import PyPDF2
import docx
import mammoth
from datetime import datetime
from app.services.ai_client import ai_client
from app.services.quick_extractor import quick_extractor
from app.utils.validators import validate_extracted_data
from app.config import settings

logger = logging.getLogger(__name__)

class PhasedCVProcessor:
    """Enhanced CV processor with phased processing capabilities"""
    
    def __init__(self):
        self.supported_mime_types = settings.allowed_file_types
        self.max_file_size = settings.max_file_size
    
    async def process_cv_phased(
        self, 
        file_content: bytes, 
        filename: str, 
        mime_type: str,
        consultant_id: str,
        phase: str = 'quick_extract',
        extracted_text: str = None,
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """
        Main phased CV processing function
        
        Phases:
        1. 'quick_extract' - Fast text extraction + basic info (2-5 seconds)
        2. 'ai_analysis' - Full AI analysis using extracted text (2+ minutes)
        """
        
        start_time = time.time()
        
        try:
            if phase == 'quick_extract':
                return await self._phase_quick_extract(
                    file_content, filename, mime_type, consultant_id, progress_callback
                )
            elif phase == 'ai_analysis':
                if not extracted_text:
                    raise Exception("Extracted text required for AI analysis phase")
                return await self._phase_ai_analysis(
                    extracted_text, filename, consultant_id, progress_callback
                )
            else:
                raise Exception(f"Unknown processing phase: {phase}")
                
        except Exception as e:
            total_time = time.time() - start_time
            logger.error(f"❌ CV processing failed in phase '{phase}' after {total_time:.2f}s: {e}")
            if progress_callback:
                progress_callback("error", 0, f"Error in {phase}: {str(e)}")
            raise Exception(f"CV processing failed in phase '{phase}': {str(e)}")
    
    async def _phase_quick_extract(
        self,
        file_content: bytes,
        filename: str,
        mime_type: str,
        consultant_id: str,
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """Phase 1: Quick text extraction and basic info extraction"""
        
        start_time = time.time()
        logger.info(f"🚀 Phase 1 - Quick extraction starting: {filename}")
        
        if progress_callback:
            progress_callback("quick_extract", 10, "Starting quick extraction...")
        
        # Step 1: Extract text from file (fast)
        if progress_callback:
            progress_callback("quick_extract", 20, "Extracting text from file...")
        
        extraction_start = time.time()
        extracted_text = await self._extract_text_from_file(
            file_content, filename, mime_type, progress_callback
        )
        extraction_time = time.time() - extraction_start
        
        if not extracted_text or len(extracted_text.strip()) < 50:
            raise Exception("Could not extract sufficient text from CV file")
        
        logger.info(f"✅ Text extracted in {extraction_time:.2f}s: {len(extracted_text)} characters")
        
        if progress_callback:
            progress_callback("quick_extract", 60, "Analyzing basic information...")
        
        # Step 2: Quick basic info extraction (rule-based)
        basic_start = time.time()
        basic_info = await quick_extractor.extract_basic_info(extracted_text, filename)
        basic_time = time.time() - basic_start
        
        logger.info(f"✅ Basic info extracted in {basic_time:.2f}s")
        logger.info(f"📋 Found: {basic_info['name']} | {basic_info['email']} | {basic_info['basic_experience']['total_years']} years")
        
        if progress_callback:
            progress_callback("quick_extract", 100, "Quick extraction completed")
        
        total_time = time.time() - start_time
        
        # Prepare response for Phase 1
        result = {
            'phase': 'quick_extract',
            'status': 'partially_processed',
            'extracted_text': extracted_text,  # Store for Phase 2
            'basic_info': basic_info,
            'consultant_updates': {
                'name': basic_info['name'],
                'email': basic_info['email'],
                'location': basic_info.get('location'),
                'processing_phase': 'partially_processed',
                'basic_info': basic_info,
                'extracted_text': extracted_text,
                'phase_timestamps': {
                    'quick_extract_started': datetime.now().isoformat(),
                    'quick_extract_completed': datetime.now().isoformat()
                },
                'extraction_metadata': {
                    'phase_1_time': total_time,
                    'text_extraction_time': extraction_time,
                    'basic_extraction_time': basic_time,
                    'file_size': len(file_content),
                    'text_length': len(extracted_text),
                    'filename': filename,
                    'mime_type': mime_type
                }
            },
            'processing_metadata': {
                'phase_1_completed': True,
                'total_time': total_time,
                'ready_for_phase_2': True
            }
        }
        
        logger.info(f"✅ Phase 1 completed in {total_time:.2f}s - Ready for AI analysis")
        return result
    
    async def _phase_ai_analysis(
        self,
        extracted_text: str,
        filename: str,
        consultant_id: str,
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """Phase 2: Full AI analysis using Qwen 2.5"""
        
        start_time = time.time()
        logger.info(f"🤖 Phase 2 - AI analysis starting for consultant {consultant_id[:8]}...")
        
        if progress_callback:
            progress_callback("ai_analysis", 10, "Starting AI analysis with Qwen 2.5...")
        
        # Perform full AI analysis
        ai_start = time.time()
        ai_analysis = await ai_client.analyze_cv_text(extracted_text, progress_callback)
        ai_time = time.time() - ai_start
        
        logger.info(f"🤖 AI analysis completed in {ai_time:.2f}s")
        logger.info(f"   Enhanced name: {ai_analysis.get('name', 'Unknown')}")
        logger.info(f"   Total experience: {ai_analysis.get('experience_years', {}).get('total', 0)} years")
        
        if progress_callback:
            progress_callback("ai_analysis", 90, "Validating AI results...")
        
        # Validate AI results
        validation_start = time.time()
        validated_data = validate_extracted_data(ai_analysis)
        validation_time = time.time() - validation_start
        
        total_time = time.time() - start_time
        
        # Prepare response for Phase 2
        result = {
            'phase': 'ai_analysis',
            'status': 'completed',
            'full_analysis': validated_data,
            'consultant_updates': {
                'name': validated_data.get('name', 'Unknown'),
                'email': validated_data.get('email'),
                'location': validated_data.get('location'),
                'experience_years': validated_data.get('experience_years'),
                'qualifications': validated_data.get('qualifications'),
                'skills': validated_data.get('skills'),
                'processing_status': 'completed',
                'processing_phase': 'completed',
                'full_analysis': validated_data,
                'phase_timestamps': {
                    'ai_analysis_started': datetime.now().isoformat(),
                    'ai_analysis_completed': datetime.now().isoformat(),
                    'fully_completed': datetime.now().isoformat()
                },
                'processing_metadata': {
                    'phase_2_time': total_time,
                    'ai_analysis_time': ai_time,
                    'validation_time': validation_time,
                    'model': settings.get_ai_config()["model"],
                    'provider': settings.ai_provider,
                    'confidence_score': ai_analysis.get('_metadata', {}).get('confidence_score', 0.0)
                }
            },
            'processing_metadata': {
                'phase_2_completed': True,
                'total_time': total_time,
                'ai_model': settings.get_ai_config()["model"]
            }
        }
        
        if progress_callback:
            progress_callback("ai_analysis", 100, "AI analysis completed")
        
        logger.info(f"✅ Phase 2 completed in {total_time:.2f}s - Consultant fully processed")
        return result
    
    async def resume_processing(
        self,
        consultant_data: Dict[str, Any],
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """Resume processing from any phase based on current state"""
        
        consultant_id = consultant_data.get('id', 'unknown')
        current_phase = consultant_data.get('processing_phase', 'pending')
        extracted_text = consultant_data.get('extracted_text')
        
        logger.info(f"🔄 Resuming processing for {consultant_id[:8]}... from phase: {current_phase}")
        
        if current_phase == 'pending':
            # Need to start from Phase 1 - but we need file content
            raise Exception("Cannot resume from pending without file content. Need to restart from upload.")
        
        elif current_phase in ['extracting', 'partially_processed']:
            if not extracted_text:
                raise Exception("Cannot resume AI analysis without extracted text")
            
            # Resume with Phase 2 (AI analysis)
            return await self._phase_ai_analysis(
                extracted_text=extracted_text,
                filename=consultant_data.get('name', 'unknown.pdf'),
                consultant_id=consultant_id,
                progress_callback=progress_callback
            )
        
        elif current_phase == 'analyzing':
            if not extracted_text:
                raise Exception("Cannot resume AI analysis without extracted text")
            
            # Retry Phase 2 (AI analysis)
            logger.info(f"🔄 Retrying AI analysis for {consultant_id[:8]}...")
            return await self._phase_ai_analysis(
                extracted_text=extracted_text,
                filename=consultant_data.get('name', 'unknown.pdf'),
                consultant_id=consultant_id,
                progress_callback=progress_callback
            )
        
        elif current_phase == 'completed':
            logger.info(f"✅ Consultant {consultant_id[:8]} already completed")
            return {
                'phase': 'completed',
                'status': 'completed',
                'message': 'Already completed',
                'consultant_updates': {},
                'processing_metadata': {'already_completed': True}
            }
        
        else:
            raise Exception(f"Unknown processing phase: {current_phase}")
    
    # Keep existing text extraction methods unchanged
    async def _extract_text_from_file(
        self, 
        file_content: bytes, 
        filename: str, 
        mime_type: str,
        progress_callback: Optional[Callable] = None
    ) -> str:
        """Extract text from different file formats"""
        
        try:
            if progress_callback:
                progress_callback("extraction", 15, f"Processing {mime_type} file...")
            
            if mime_type == 'application/pdf':
                text = await self._extract_from_pdf(file_content, progress_callback)
            elif mime_type in [
                'application/msword', 
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            ]:
                text = await self._extract_from_word(file_content, mime_type, progress_callback)
            else:
                raise Exception(f"Unsupported file type: {mime_type}")
            
            # Clean and validate extracted text
            cleaned_text = self._clean_extracted_text(text)
            
            if len(cleaned_text) < 50:
                raise Exception("File appears to be empty or contains mostly images/scanned content")
            
            return cleaned_text
            
        except Exception as e:
            logger.error(f"❌ Text extraction failed for {filename}: {e}")
            raise Exception(f"Failed to extract text from {filename}: {str(e)}")
    
    async def _extract_from_pdf(
        self, 
        file_content: bytes, 
        progress_callback: Optional[Callable] = None
    ) -> str:
        """Extract text from PDF with improved error handling"""
        
        try:
            if progress_callback:
                progress_callback("extraction", 20, "Reading PDF content...")
            
            file_stream = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(file_stream)
            
            if len(pdf_reader.pages) == 0:
                raise Exception("PDF has no pages")
            
            text_parts = []
            total_pages = len(pdf_reader.pages)
            successful_pages = 0
            
            for i, page in enumerate(pdf_reader.pages):
                if progress_callback:
                    progress = 20 + (i / total_pages) * 10  # 20% to 30%
                    progress_callback("extraction", progress, f"Extracting page {i+1}/{total_pages}...")
                
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)
                        successful_pages += 1
                except Exception as page_error:
                    logger.warning(f"⚠️ Could not extract text from page {i+1}: {page_error}")
                    continue
            
            if not text_parts:
                raise Exception("Could not extract text from any PDF pages")
            
            full_text = '\n'.join(text_parts)
            logger.info(f"✅ PDF extraction: {len(full_text)} chars from {successful_pages}/{total_pages} pages")
            
            return full_text
            
        except Exception as e:
            raise Exception(f"PDF processing failed: {str(e)}")
    
    async def _extract_from_word(
        self, 
        file_content: bytes, 
        mime_type: str, 
        progress_callback: Optional[Callable] = None
    ) -> str:
        """Extract text from Word documents"""
        
        try:
            if progress_callback:
                progress_callback("extraction", 20, "Reading Word document...")
            
            if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # DOCX file
                text = await self._extract_from_docx(file_content, progress_callback)
            else:
                # DOC file - use mammoth
                text = await self._extract_from_doc(file_content, progress_callback)
            
            if len(text.strip()) < 50:
                raise Exception("Word document appears to be empty")
            
            logger.info(f"✅ Word extraction: {len(text)} characters")
            return text
            
        except Exception as e:
            raise Exception(f"Word document processing failed: {str(e)}")
    
    async def _extract_from_docx(
        self, 
        file_content: bytes, 
        progress_callback: Optional[Callable] = None
    ) -> str:
        """Extract text from DOCX files"""
        
        file_stream = io.BytesIO(file_content)
        doc = docx.Document(file_stream)
        
        text_parts = []
        total_paragraphs = len(doc.paragraphs)
        
        # Extract paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            if progress_callback and i % 20 == 0:  # Update every 20 paragraphs
                progress = 20 + (i / total_paragraphs) * 8
                progress_callback("extraction", progress, f"Processing paragraph {i+1}/{total_paragraphs}...")
            
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract tables
        if progress_callback:
            progress_callback("extraction", 28, "Extracting tables...")
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        return '\n'.join(text_parts)
    
    async def _extract_from_doc(
        self, 
        file_content: bytes, 
        progress_callback: Optional[Callable] = None
    ) -> str:
        """Extract text from DOC files using mammoth"""
        
        if progress_callback:
            progress_callback("extraction", 25, "Converting DOC file...")
        
        file_stream = io.BytesIO(file_content)
        result = mammoth.extract_raw_text(file_stream)
        
        # Log any warnings from mammoth
        if result.messages:
            for message in result.messages:
                logger.warning(f"⚠️ Mammoth warning: {message}")
        
        return result.value
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text for better AI processing"""
        
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive newlines but preserve structure
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Clean up common PDF extraction artifacts
        text = re.sub(r'[^\w\s\n\-\.,;:@/()&%$#!?+="\']+', ' ', text)
        
        # Remove very short lines that are likely artifacts
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if len(line) >= 3:  # Keep lines with at least 3 characters
                cleaned_lines.append(line)
        
        # Rejoin and clean final whitespace
        text = '\n'.join(cleaned_lines)
        text = re.sub(r' {2,}', ' ', text)  # Multiple spaces to single space
        
        return text.strip()

# Global phased CV processor instance
cv_processor = PhasedCVProcessor()